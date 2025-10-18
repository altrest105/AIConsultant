import logging
import time
import numpy as np
import torch
from tqdm import tqdm
from docx import Document as DocxDocument
from sentence_transformers import CrossEncoder
from rank_bm25 import BM25Okapi
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import Qdrant
from langchain_core.documents import Document
from qdrant_client import QdrantClient
from qdrant_client.models import Distance, VectorParams
from django.conf import settings

logger = logging.getLogger(__name__)

# Глобальные переменные и конфигурация
VECTOR_STORE = None
EMBEDDINGS_MODEL = None
RERANKER = None

BM25_INDEX = None
BM25_CORPUS = []
FULL_INDEXED_DOCUMENTS = []

HEADER_PRIORITY = {'H0': 6.0, 'H1': 5.0, 'H2': 4.0, 'H3': 3.0, 'H4': 2.0, 'T': 1.0, 'L': 1.0}

# Инициализация всех компонентов QA
def initialize_embeddings():
    """Инициализирует и загружает модель для создания векторных представлений (embeddings)."""
    global EMBEDDINGS_MODEL

    embedding_model_name = settings.QA_CONFIG.get("EMBEDDING_MODEL_NAME")
    device = 'cuda' if torch.cuda.is_available() else 'cpu'

    logger.info(f"🔄 Загрузка Embedding: {embedding_model_name}")
    try:
        EMBEDDINGS_MODEL = HuggingFaceEmbeddings(
            model_name=embedding_model_name,
            model_kwargs={'device': device}
        )
        logger.info(f"✅ Embedding модель загружена на: {device}")
    except Exception as e:
        logger.error(f"❌ Ошибка при загрузке Embedding: {e}")
        EMBEDDINGS_MODEL = None

def initialize_reranker():
    """Инициализирует и загружает модель для переранжирования (reranker)."""
    global RERANKER

    reranker_model_name = settings.QA_CONFIG.get("RERANKER_MODEL_NAME", "BAAI/bge-reranker-v2-m3")
    reranker_max_length = settings.QA_CONFIG.get("RERANKER_MAX_LENGTH", 512)
    device = 'cuda' if torch.cuda.is_available() else 'cpu'

    logger.info(f"🔄 Загрузка Reranker: {reranker_model_name}")
    try:
        RERANKER = CrossEncoder(
            reranker_model_name,
            max_length=reranker_max_length,
            device=device
        )
        logger.info(f"✅ Reranker загружен на: {device}")
    except Exception as e:
        logger.error(f"❌ Ошибка при загрузке Reranker: {e}")
        RERANKER = None

def initialize_vector_store():
    """Инициализирует клиент для векторного хранилища Qdrant."""
    global VECTOR_STORE, EMBEDDINGS_MODEL
    logger.info("🔄 Инициализация Qdrant...")
    if not EMBEDDINGS_MODEL:
        logger.error("❌ EMBEDDINGS_MODEL не инициализирован.")
        return
    try:
        qdrant_path = settings.QA_CONFIG.get("QDRANT_PATH")
        qdrant_name = settings.QA_CONFIG.get("QDRANT_NAME")
        client = QdrantClient(path=qdrant_path)

        VECTOR_STORE = Qdrant(
            client=client, 
            collection_name=qdrant_name, 
            embeddings=EMBEDDINGS_MODEL
        )
        logger.info("✅ Qdrant инициализирован.")
    except Exception as e:
        logger.error(f"❌ Ошибка при инициализации Qdrant: {e}")
        VECTOR_STORE = None

def initialize_qa_system():
    """Выполняет полную инициализацию всех компонентов QA-системы."""
    global EMBEDDINGS_MODEL, RERANKER, VECTOR_STORE
    
    if not (EMBEDDINGS_MODEL and RERANKER and VECTOR_STORE):
        logger.info(f"🔄 Инициализация системы QA")
        if not EMBEDDINGS_MODEL: initialize_embeddings()
        if not RERANKER: initialize_reranker()
        if EMBEDDINGS_MODEL and not VECTOR_STORE: initialize_vector_store()
        logger.info("✅ QA система инициализирована.")

# Парсинг DOCX
def is_bold_paragraph(p):
    """Проверяет, является ли параграф в DOCX-документе жирным."""
    text = p.text.strip()
    if not text:
        return False

    for run in p.runs:
        if run.text.strip(): 
            return run.bold or ('strong' in run.element.xml) 
    return False

def equals_indent(indent_emu, target_cm, eps_cm=0.01):
    """Сравнивает отступ параграфа с целевым значением в сантиметрах."""
    return abs(indent_emu.cm - target_cm) < eps_cm

def get_chunk_type(paragraph):
    """Определяет семантический тип параграфа (заголовок, текст, список)."""
    style_name = paragraph.style.name.lower()
    text = paragraph.text.strip()
    if not text: return None
    
    # Проверка на формальные стили (H0, H1, H2)
    if style_name.startswith('heading 1') or 'заголовок 1' in style_name:
        return 'H0'
    if style_name.startswith('heading 2') or 'заголовок 2' in style_name:
        return 'H1'
    if style_name.startswith('heading 3') or 'заголовок 3' in style_name:
        return 'H2'
        
    # Логика для H3 и H4 ("Обычный текст" + Жирность + Отступ)
    is_normal_style = 'normal' in style_name or 'обычный текст' in style_name
    
    if is_normal_style and is_bold_paragraph(paragraph):

        first_line_indent_style = paragraph.style.paragraph_format.first_line_indent
        first_line_indent = paragraph.paragraph_format.first_line_indent

        if first_line_indent is not None and equals_indent(first_line_indent, 1.25):
            return 'H4'
        
        if first_line_indent is None and equals_indent(first_line_indent_style, 0.75):
            return 'H3'
        
    # Проверка на списки L
    if 'list' in style_name or 'bullet' in style_name or 'number' in style_name:
        return 'L'

    # Проверка на цитаты, обычный текст T
    if 'quote' in style_name or 'normal' in style_name:
        return None
        
    return None

def parse_document_universal(file_path):
    """Парсит DOCX-файл, разбивая его на семантические чанки (заголовки, текст)."""
    logger.info(f"📄 Парсинг файла: {file_path.name}")
    
    if file_path.suffix.lstrip('.') != 'docx':
        logger.warning(f"⚠️ Пропущен файл: {file_path.name}. Поддерживается только DOCX.")
        return []
         
    documents = []
    paragraphs_data = [p for p in DocxDocument(file_path).paragraphs if p.text.strip()]
        
    i = 0
    paragraphs_data_count = len(paragraphs_data)
    
    current_header_h0: str = "" 
    current_header_h1: str = "" 
    current_header_h2: str = ""
    current_header_h3: str = "" 
    current_header_h4: str = ""
    
    while i < paragraphs_data_count:
        
        current_paragraph_obj = paragraphs_data[i]
        current_paragraph = current_paragraph_obj.text.strip()
        header_level = get_chunk_type(current_paragraph_obj)
        
        # Обновляем заголовки, если текущий абзац — заголовок
        if header_level and header_level != 'L':
            
            # Логика заголовка H0-H4)
            if header_level == 'H0':
                current_header_h0 = current_paragraph
                current_header_h1 = current_header_h2 = current_header_h3 = current_header_h4 = ""
            elif header_level == 'H1':
                current_header_h1 = current_paragraph
                current_header_h2 = current_header_h3 = current_header_h4 = ""
            elif header_level == 'H2':
                current_header_h2 = current_paragraph
                current_header_h3 = current_header_h4 = ""
            elif header_level == 'H3':
                current_header_h3 = current_paragraph
                current_header_h4 = ""
            elif header_level == 'H4':
                current_header_h4 = current_paragraph
            
            header_context_parts = [h for h in [current_header_h0, current_header_h1, current_header_h2, current_header_h3, current_header_h4] if h]
            header_context = " | ".join(header_context_parts)
            header_chunk_content = f"[{header_context}] {current_paragraph}"
            
            documents.append(
                Document(
                    page_content=header_chunk_content,
                    metadata={
                        "source": file_path.name,
                        "chunk_type": header_level, 
                        "h0_header": current_header_h0,
                        "h1_header": current_header_h1,
                        "h2_header": current_header_h2,
                        "h3_header": current_header_h3,
                        "h4_header": current_header_h4,
                        "start_index": i,            
                        "end_index": i,              
                        "original_text": current_paragraph
                    }
                )
            )
            i += 1 
            continue
            
        # Объединение T и L
        current_chunk_raw_texts = [current_paragraph]
        start_paragraph_index = i
        j = i
        final_chunk_type = header_level if header_level is not None else 'T'
        
        ends_with_colon = current_paragraph.rstrip().endswith(':')
        should_aggregate = False

        # Агрегация текст ":" + Список "L"
        if final_chunk_type == 'T' and ends_with_colon and (j + 1 < paragraphs_data_count):
            next_paragraph_type = get_chunk_type(paragraphs_data[j + 1])
            if next_paragraph_type == 'L':
                should_aggregate = True
        
        # Цикл агрегации
        if should_aggregate:
            final_chunk_type = 'T' 

            while j + 1 < paragraphs_data_count:
                next_paragraph_obj = paragraphs_data[j + 1]
                next_paragraph_text = next_paragraph_obj.text.strip()
                
                if not next_paragraph_text:
                    j += 1
                    continue

                next_header_level = get_chunk_type(next_paragraph_obj) 

                if next_header_level != 'L':
                    break
                    
                current_chunk_raw_texts.append(next_paragraph_text)
                j += 1
        
        current_chunk_raw = "\n".join(current_chunk_raw_texts)
        
        header_context_parts = [h for h in [current_header_h0, current_header_h1, current_header_h2, current_header_h3, current_header_h4] if h]
        header_context = " | ".join(header_context_parts)
        final_chunk_content = f"[{header_context}] {current_chunk_raw}"

        documents.append(
            Document(
                page_content=final_chunk_content,
                metadata={
                    "source": file_path.name,
                    "chunk_type": final_chunk_type, 
                    "h0_header": current_header_h0,
                    "h1_header": current_header_h1,
                    "h2_header": current_header_h2,
                    "h3_header": current_header_h3,
                    "h4_header": current_header_h4,
                    "start_index": start_paragraph_index,
                    "end_index": j,
                    "original_text": current_chunk_raw
                }
            )
        )
        
        # Передвигаем основной индекс на следующий абзац после объединенных
        i = j + 1 

    logger.info(f"📚 Разделено на {len(documents)} чанков.")
    return documents


# Обновление BM25 индекса
def update_bm25_index(documents):
    """Обновляет BM25-индекс и глобальный список документов."""
    global BM25_INDEX, BM25_CORPUS, FULL_INDEXED_DOCUMENTS

    texts = [doc.page_content for doc in documents]

    if not FULL_INDEXED_DOCUMENTS or len(FULL_INDEXED_DOCUMENTS) == 0:
         BM25_CORPUS = []
         FULL_INDEXED_DOCUMENTS = []
    
    BM25_CORPUS.extend(texts)
    FULL_INDEXED_DOCUMENTS.extend(documents)
    tokenized_corpus = [doc.split(" ") for doc in BM25_CORPUS]
    BM25_INDEX = BM25Okapi(tokenized_corpus)
    logger.info(f"✅ Индексы (Qdrant, BM25) и структура (FULL_INDEXED_DOCUMENTS) обновлены.")

def chunk_list(data, size):
    """Разделяет список на батчи заданного размера."""
    for i in range(0, len(data), size):
        yield data[i:i + size]

def index_document(file_path):
    """Индексирует один документ, добавляя его чанки в Qdrant и BM25."""
    global VECTOR_STORE

    documents = parse_document_universal(file_path)
    BATCH_SIZE = settings.QA_CONFIG.get("QDRANT_BATCH_SIZE", 32)

    if not documents: return
    if VECTOR_STORE:
        try:
            logger.info(f"🔄 Qdrant: Начинаем загрузку {len(documents)} чанков (батч: {BATCH_SIZE}).")
            chunk_batches = list(chunk_list(documents, BATCH_SIZE))
            
            for batch in tqdm(chunk_batches, desc=f"Qdrant: {file_path.name}", unit="batch"):
                VECTOR_STORE.add_documents(batch) 
            
            logger.info(f"✅ Qdrant: Добавлено {len(documents)} чанков.")
        except Exception as e:
             logger.error(f"❌ Ошибка Qdrant при добавлении векторов: {e}")
    update_bm25_index(documents)

def index_knowledge_base(folder_path):
    """Полностью переиндексирует базу знаний из указанной папки."""
    logger.info(f"📚 Индексация базы знаний")
    if not folder_path.exists():
        logger.error(f"❌ Папка не найдена: {folder_path}")
        return
    
    supported_formats = ['.docx'] 
    files_to_index = []

    for ext in supported_formats:
        files_to_index.extend([f for f in folder_path.glob(f"*{ext}") if not f.name.startswith('~$')])
    
    if not files_to_index:
        logger.warning(f"⚠️ Нет файлов DOCX для индексации в {folder_path}")
        return
    
    if VECTOR_STORE:
        VECTOR_STORE.client.recreate_collection(
            collection_name=VECTOR_STORE.collection_name,
            vectors_config=VectorParams(size=1024, distance=Distance.COSINE)
        )
        logger.info(f"🗑️ Qdrant: Коллекция '{VECTOR_STORE.collection_name}' очищена.")
        
    global BM25_CORPUS, BM25_INDEX, FULL_INDEXED_DOCUMENTS
    BM25_CORPUS = []
    BM25_INDEX = None
    FULL_INDEXED_DOCUMENTS = []

    for file_path in tqdm(files_to_index, desc="Индексация"):
        index_document(file_path)

    logger.info(f"✅ Индексация завершена.")
    logger.info(f"    Всего чанков в корпусе: {len(FULL_INDEXED_DOCUMENTS)}")

# Ретивер
class CustomRetriever:
    """Гибридный ретривер, использующий векторный поиск, BM25 и реранжирование."""
    def __init__(self, vector_store, reranker):
        self.vector_store = vector_store
        self.reranker = reranker

    def get_relevant_documents(self, query):
        """Находит и ранжирует наиболее релевантные документы для заданного вопроса."""
        global BM25_INDEX, BM25_CORPUS
        
        vec_results = self.vector_store.similarity_search_with_score(query, k=settings.QA_CONFIG.get("K_VEC", 40))
        bm25_results = []

        if BM25_INDEX and BM25_CORPUS:
            tokenized_query = query.split(" ")
            doc_scores = BM25_INDEX.get_scores(tokenized_query)
            top_n_indices = np.argsort(doc_scores)[::-1][:settings.QA_CONFIG.get("K_BM25", 40)]
            
            for idx in top_n_indices:
                doc = FULL_INDEXED_DOCUMENTS[idx]
                bm25_results.append(doc)

        all_docs_map = {}
        K_RRF = 60 
        def calculate_rrf_score(rank, k=K_RRF): return 1 / (k + rank)

        for rank, (doc, _) in enumerate(vec_results):
            text_hash = hash(doc.page_content)
            if text_hash not in all_docs_map: all_docs_map[text_hash] = {"doc": doc, "rrf_score": 0.0}
            all_docs_map[text_hash]["rrf_score"] += calculate_rrf_score(rank + 1) * settings.QA_CONFIG.get("ALPHA", 0.5)

        for rank, doc in enumerate(bm25_results):
            text_hash = hash(doc.page_content)
            if text_hash not in all_docs_map: all_docs_map[text_hash] = {"doc": doc, "rrf_score": 0.0}
            all_docs_map[text_hash]["rrf_score"] += calculate_rrf_score(rank + 1) * (1.0 - settings.QA_CONFIG.get("ALPHA", 0.5))

        final_candidates = sorted(all_docs_map.values(), key=lambda x: x["rrf_score"], reverse=True)
        rerank_input = [item["doc"] for item in final_candidates[:50]]
        
        if self.reranker and len(rerank_input) > 0:
            pairs = [[query, doc.page_content] for doc in rerank_input]
            scores = self.reranker.predict(pairs)
            probabilities = 1 / (1 + np.exp(-scores))
            
            scored_documents = sorted(
                zip(rerank_input, probabilities),
                key=lambda x: x[1],
                reverse=True
            )
            return scored_documents[:settings.QA_CONFIG.get("TOP_N_RERANKER", 5)]
        
        return []

# Расширение контекста и ответ
def expand_context(best_match_doc):
    """Расширяет контекст для найденного заголовка, собирая последующий текст."""
    global FULL_INDEXED_DOCUMENTS, HEADER_PRIORITY
    
    chunk_type = best_match_doc.metadata.get('chunk_type')
    original_text_target = best_match_doc.metadata.get('original_text', best_match_doc.page_content)
    
    # Если это текст, возвращаем его без расширения
    if chunk_type == 'T':
        return original_text_target
        
    target_level = chunk_type 
    target_priority = HEADER_PRIORITY.get(target_level, 0)
    
    # Получаем индекс из метаданных
    start_index_in_full = best_match_doc.metadata.get('start_index', -1) 
    
    if start_index_in_full == -1:
        logger.error(f"❌ Ошибка расширения: Не найден start_index в метаданных документа. Возврат только заголовка.")
        return original_text_target
        
    expanded_text = []
    
    # Добавляем сам заголовок
    expanded_text.append(f"<{target_level}> {original_text_target} </{target_level}>")
    
    # Начинаем расширение со следующего индекса
    for idx in range(start_index_in_full + 1, len(FULL_INDEXED_DOCUMENTS)):
        current_doc = FULL_INDEXED_DOCUMENTS[idx]
        current_type = current_doc.metadata.get('chunk_type')
        current_priority = HEADER_PRIORITY.get(current_type, 0)
        current_original_text = current_doc.metadata.get('original_text', '')
        
        # Условие остановки: Заголовок того же или более высокого уровня
        if current_type != 'T' and current_priority >= target_priority:
            break
        
        # Добавляем контент (Текст или Заголовок более низкого уровня)
        if current_type == 'T':
            expanded_text.append(current_original_text)
        else:
            expanded_text.append(f"<{current_type}> {current_original_text} </{current_type}>")
            
    logger.info(f"✅ Расширение завершено. Собрано {len(expanded_text)} элементов.")
    return "\n\n".join(expanded_text)

def answer_question(question):
    """Основная функция для ответа на вопрос, использующая всю цепочку RAG."""
    global VECTOR_STORE, RERANKER, FULL_INDEXED_DOCUMENTS, HEADER_PRIORITY
    
    if not (VECTOR_STORE and RERANKER and FULL_INDEXED_DOCUMENTS):
        logger.error("❌ Система поиска/индексации не инициализирована.")
        return {"answer": "Система поиска не инициализирована.", "source_documents": [], "confidence": 0.0, "latency_sec": 0.0}

    start_time = time.time()
    
    retriever = CustomRetriever(vector_store=VECTOR_STORE, reranker=RERANKER)
    scored_documents = retriever.get_relevant_documents(question)
    
    latency = time.time() - start_time
    
    if not scored_documents:
        return {"answer": "К сожалению, в базе знаний не найдено информации, релевантной вашему запросу. Пожалуйста, попробуйте переформулировать вопрос.", "source_documents": [], "confidence": 0.0, "latency_sec": latency}

    best_match_doc = scored_documents[0][0]
    top_score = scored_documents[0][1]

    normalized_query = question.strip().lower()
    priority_match = None
    max_priority = -1
    
    # Логика Приоритетного выбора
    for doc, score in scored_documents[:settings.QA_CONFIG.get("TOP_N_RERANKER", 5)]:
        doc_type = doc.metadata.get('chunk_type')
        doc_original_text = doc.metadata.get('original_text', '')
        
        if doc_type != 'T' and doc_original_text.strip().lower() == normalized_query and score > 0.8:
            priority = HEADER_PRIORITY.get(doc_type, 0)
            
            if priority > max_priority:
                max_priority = priority
                priority_match = doc

    if priority_match:
        best_match_doc = priority_match
        top_score = scored_documents[0][1] 
    
    # Расширение контекста
    final_expanded_text = expand_context(best_match_doc)
    
    
    answer_text = f"{final_expanded_text}"
    answer_type = best_match_doc.metadata.get('chunk_type')
    
    logger.info(f"⏱️ Ответ дан за: {latency:.2f}s")
    
    if top_score < settings.QA_CONFIG.get("CONFIDENCE_THRESHOLD", 0.55):
        logger.warning(f"⚠️ Низкий уровень доверия: {top_score:.2f}")
        return {
            "answer": "К сожалению, в базе знаний не найдено информации, релевантной вашему запросу. Пожалуйста, попробуйте переформулировать вопрос.",
            "source_documents": [best_match_doc],
            "confidence": top_score,
            "latency_sec": latency,
            "chunk_type": answer_type
        }

    return {
        "answer": answer_text,
        "source_documents": [best_match_doc],
        "confidence": top_score,
        "latency_sec": latency,
        "chunk_type": answer_type
    }
